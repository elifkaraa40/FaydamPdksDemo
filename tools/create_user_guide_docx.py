from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

ROOT = Path(__file__).resolve().parents[1]
SHOTS = ROOT / "tmp" / "guide-screens"
OUT = ROOT / "docs" / "Faydam_PDKS_Kullanim_ve_Yardim_Dokumani.docx"

NAVY = "12213A"
BLUE = "2563EB"
PURPLE = "5B4FE9"
CYAN = "22A6C7"
MUTED = "64748B"
LIGHT = "F1F5F9"
PALE_BLUE = "EAF2FF"
GREEN = "0F9F6E"
RED = "C24155"
WHITE = "FFFFFF"


def set_font(run, size=None, bold=None, color=None, name="Calibri", italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=140, bottom=110, end=140):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    table.autofit = False
    total = sum(widths_dxa)
    tblpr = table._tbl.tblPr
    tblw = tblpr.first_child_found_in("w:tblW")
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    ind = tblpr.first_child_found_in("w:tblInd")
    if ind is None:
        ind = OxmlElement("w:tblInd")
        tblpr.append(ind)
    ind.set(qn("w:w"), "120")
    ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths_dxa):
            tcw = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tcw.set(qn("w:w"), str(width))
            tcw.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sayfa ")
    set_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def configure(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    specs = {
        "Heading 1": (18, NAVY, 16, 8),
        "Heading 2": (14, PURPLE, 12, 6),
        "Heading 3": (11.5, NAVY, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.18

    if "Figure Caption" not in [s.name for s in doc.styles]:
        cap = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = "Calibri"
        cap.font.size = Pt(8.5)
        cap.font.italic = True
        cap.font.color.rgb = RGBColor.from_string(MUTED)
        cap.paragraph_format.space_before = Pt(3)
        cap.paragraph_format.space_after = Pt(8)
        cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("FAYDAM PDKS  |  KULLANIM VE YARDIM"), size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def add_kicker(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text.upper())
    set_font(r, size=8.5, bold=True, color=PURPLE)
    return p


def add_callout(doc, title, body, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(title), bold=True, color=accent, size=10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    set_font(p2.add_run(body), size=10, color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_steps(doc, steps):
    numbering = doc.part.numbering_part.element
    style_num_id = doc.styles["List Number"]._element.pPr.numPr.numId.val
    base_num = next(n for n in numbering.findall(qn("w:num")) if int(n.get(qn("w:numId"))) == style_num_id)
    abstract_id = base_num.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    new_num_id = max(existing_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        num_pr = p._p.get_or_add_pPr().get_or_add_numPr()
        num_pr.get_or_add_ilvl().val = 0
        num_pr.get_or_add_numId().val = new_num_id
        p.add_run(step)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def set_picture_alt(inline_shape, description):
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def add_image(doc, filename, caption, width=6.25, max_height=5.35):
    path = SHOTS / filename
    with Image.open(path) as image:
        ratio = image.width / image.height
    width = min(width, max_height * ratio)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    set_picture_alt(shape, caption)
    doc.add_paragraph(caption, style="Figure Caption")


def new_page(doc, kicker, title, intro=None):
    doc.add_page_break()
    add_kicker(doc, kicker)
    doc.add_heading(title, level=1)
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(8)


def add_cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(60)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("KULLANIM VE YARDIM DOKÜMANI")
    set_font(r, size=10, bold=True, color=PURPLE)
    p.paragraph_format.space_after = Pt(16)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Faydam PDKS")
    set_font(r, size=32, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(7)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Personel devam, vardiya, izin ve puantaj yönetimi"), size=14, color=MUTED)
    p.paragraph_format.space_after = Pt(18)
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_shape = logo_p.add_run().add_picture(str(ROOT / "FaydamPDKS.Web" / "wwwroot" / "images" / "faydam-navbar-logo.png"), width=Inches(2.35))
    set_picture_alt(logo_shape, "Faydam PDKS logosu")
    logo_p.paragraph_format.space_after = Pt(42)
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [2700, 6660])
    rows = [
        ("Kapsam", "Web yönetim paneli ve personel portalı"),
        ("Hedef kullanıcı", "Yönetici, insan kaynakları yetkilisi ve personel"),
        ("Sürüm", "1.0"),
        ("Güncelleme", "22 Temmuz 2026"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        shade_cell(row.cells[0], PALE_BLUE)
        set_font(row.cells[0].paragraphs[0].add_run(label), bold=True, color=PURPLE)
        set_font(row.cells[1].paragraphs[0].add_run(value), color=NAVY)


def add_contents(doc):
    new_page(doc, "Belge haritası", "İçindekiler", "Kılavuz, görev odaklı ilerleyecek şekilde yönetici ve personel akışlarına ayrılmıştır.")
    sections = [
        ("1", "Sisteme giriş ve genel kullanım"),
        ("2", "Yönetici kontrol paneli"),
        ("3", "PDKS yönetimi: giriş-çıkış, izin, vardiya ve takvim"),
        ("4", "Organizasyon, çalışma konumları ve personel"),
        ("5", "Raporlama, denetim ve işlemler"),
        ("6", "Hesap, arama, bildirim ve görünüm ayarları"),
        ("7", "Personel portalı"),
        ("8", "Sık karşılaşılan durumlar ve güvenli kullanım"),
    ]
    table = doc.add_table(rows=len(sections), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_geometry(table, [900, 8460])
    for row, (number, title) in zip(table.rows, sections):
        shade_cell(row.cells[0], PURPLE)
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(row.cells[0].paragraphs[0].add_run(number), bold=True, color=WHITE, size=11)
        set_font(row.cells[1].paragraphs[0].add_run(title), bold=True, color=NAVY, size=10.5)
    doc.add_paragraph()
    add_callout(doc, "Rol farkı", "Yönetici hesabı tüm yönetim modüllerini görür. Personel hesabı yalnızca kendi çalışma kayıtları, izinleri, konumları, bildirimleri ve kişisel verilerine erişir.")


def build():
    doc = Document()
    configure(doc)
    add_cover(doc)
    add_contents(doc)

    new_page(doc, "1 · Başlangıç", "Sisteme giriş", "Faydam PDKS’ye kurumunuzun size bildirdiği web adresinden erişin.")
    add_steps(doc, [
        "E-posta alanına kurumsal e-posta adresinizi yazın.",
        "Parola alanına parolanızı girin. Gerekirse göz simgesiyle parolayı geçici olarak görünür yapın.",
        "Kişisel ve güvenli bir cihaz kullanıyorsanız Oturumu açık tut seçeneğini işaretleyin.",
        "Giriş yap düğmesine basın. Rolünüze göre yönetici paneli veya personel portalı açılır.",
    ])
    add_image(doc, "01-giris.png", "Şekil 1 — Faydam PDKS giriş ekranı", 5.9)
    add_callout(doc, "Giriş yapılamıyorsa", "E-posta ve parolayı yeniden kontrol edin. Hesap pasif, onay bekliyor veya kilitli ise sistem yöneticinizle iletişime geçin.", fill="FFF7ED", accent="B45309")

    new_page(doc, "2 · Yönetici", "Kontrol Paneli", "Kontrol Paneli, ekibin günlük devam durumunu ve yöneticinin bekleyen işlerini tek ekranda özetler.")
    add_bullets(doc, [
        "İş yerinde, geç başlayan, izinli ve giriş kaydı olmayan personel sayılarını izleyin.",
        "Çalışan devam durumu grafiğiyle günlük dağılımı karşılaştırın.",
        "Son hareketler alanından en yeni giriş-çıkış kayıtlarını kontrol edin.",
        "Onay bekleyenler alanından izin ve düzeltme taleplerine geçin.",
    ])
    add_image(doc, "02-kontrol-paneli.png", "Şekil 2 — Yönetici Kontrol Paneli ve günlük özet", 6.25)

    new_page(doc, "3 · PDKS", "PDKS Yönetim Merkezi", "PDKS menüsü, devam yönetiminde kullanılan tüm tanım ve işlemlere görev odaklı erişim sağlar.")
    add_image(doc, "03-pdks-merkezi.png", "Şekil 3 — PDKS Yönetim Merkezi", 6.15)
    add_bullets(doc, [
        "Giriş-Çıkış: QR kodları ve geçiş kayıtları",
        "İzin Yönetimi: taleplerin incelenmesi ve sonuçlandırılması",
        "Vardiya ve Planlama: vardiya tanımı ve personel ataması",
        "Çalışma Takvimi: tatil, hafta tatili ve özel çalışma günleri",
        "Tanımlamalar ve Çalışma Konumları: işyeri, bölüm ve saha kuralları",
    ])

    new_page(doc, "3.1 · PDKS", "Giriş-Çıkış Yönetimi", "Mobil uygulamada kullanılacak QR kodlarını üretin; yenileme, pasife alma ve geçiş izleme işlemlerini aynı ekrandan yönetin.")
    add_steps(doc, [
        "Yeni QR alanında QR adı, işyeri, giriş-çıkış bölgesi ve işlem türünü seçin.",
        "QR oluştur düğmesiyle kodu üretin.",
        "Süresi dolan veya paylaşımı riskli kodları Yenile ile değiştirin.",
        "Kullanılmayacak kodu Pasife al seçeneğiyle devre dışı bırakın.",
        "Geçiş Kayıtları sekmesinden personel, bölge, işlem ve kaynak bilgisini inceleyin.",
    ])
    add_image(doc, "04-giris-cikis.png", "Şekil 4 — QR kodları ve geçiş kayıtları", 5.65)

    new_page(doc, "3.2 · PDKS", "İzin Yönetimi", "Yöneticiler personelin izin taleplerini tarih, tür, açıklama ve mevcut durumuyla birlikte değerlendirir.")
    add_steps(doc, [
        "Listeden değerlendirilecek talebi bulun.",
        "Tarih aralığı, izin türü ve açıklamayı kontrol edin.",
        "Uygunsa Onayla, uygun değilse Reddet işlemini seçin.",
        "Kararın personele bildirim olarak ulaştığını doğrulayın.",
    ])
    add_image(doc, "05-izin-yonetimi.png", "Şekil 5 — Yönetici izin talepleri listesi", 6.2)
    add_callout(doc, "Dikkat", "Onay veya ret kararı denetim izine kaydedilir. Karar vermeden önce tarih aralığı ile çalışma takvimini birlikte kontrol edin.", fill="FFF1F2", accent=RED)

    new_page(doc, "3.3 · PDKS", "Vardiya ve Planlama", "Vardiya tanımları çalışma başlangıcı, bitişi ve mola kurallarını; atamalar ise ilgili personel ve tarih aralığını belirler.")
    add_steps(doc, [
        "Yeni vardiya formunda vardiya adı, başlangıç ve bitiş saatini girin.",
        "Gece vardiyası gerekiyorsa ertesi güne taşan saat aralığını doğru tanımlayın.",
        "Vardiyayı kaydedin ve personel atama alanından çalışan ile tarih aralığını seçin.",
        "Atama listesinden geçerlilik tarihlerini kontrol edin.",
    ])
    add_image(doc, "06-vardiya-planlama.png", "Şekil 6 — Vardiya tanımı ve personel atamaları", 5.8)

    new_page(doc, "3.4 · PDKS", "Çalışma Takvimi", "Takvim; resmi tatil, hafta tatili ve özel çalışma günlerinin puantaj hesabına doğru yansımasını sağlar.")
    add_steps(doc, [
        "Tarih ve gün tipini seçin.",
        "Kayıt tüm organizasyon için geçerliyse kapsamı genel bırakın; yalnızca bir işyeri içinse ilgili işyerini seçin.",
        "Gerekli açıklamayı yazın ve Takvime ekle düğmesine basın.",
        "Aynı tarih için çakışan kayıt bulunmadığını kontrol edin.",
    ])
    add_image(doc, "07-calisma-takvimi.png", "Şekil 7 — Çalışma takvimi ve gün tipi tanımı", 6.0)

    new_page(doc, "4 · Organizasyon", "İşyeri ve Bölümler", "Organizasyon ekranında işyerleri ve bağlı bölümler tanımlanır. Personel, vardiya ve konum işlemleri bu yapıyı kullanır.")
    add_image(doc, "08-organizasyon.png", "Şekil 8 — İşyeri ve bölüm tanımları", 6.1)
    add_bullets(doc, [
        "İşyeri kodlarını kısa, benzersiz ve kurumsal adlandırma standardına uygun seçin.",
        "Bölümü doğru işyerine bağlayın.",
        "Kullanılmayan kayıtları silmek yerine pasif yapmayı tercih edin.",
    ])

    new_page(doc, "4.1 · Organizasyon", "Çalışma Konumları", "Ofis, saha, müşteri veya proje alanı gibi çalışma konumlarını tanımlayın ve personel taleplerini değerlendirin.")
    add_bullets(doc, [
        "Konum adı, türü ve ilgili işyerini seçin.",
        "Saha doğrulaması kullanılacaksa koordinat ve yarıçap bilgilerini dikkatle girin.",
        "Personelin konum talebini onaylamadan önce tarih, gerekçe ve atama kapsamını kontrol edin.",
    ])
    add_image(doc, "09-calisma-konumlari.png", "Şekil 9 — Çalışma konumu tanımları ve talepler", 5.8)

    new_page(doc, "4.2 · Personel", "Personel Yönetimi", "Personel ekranı; çalışan oluşturma, rol ve organizasyon atama, hesap durumu ve temel iletişim bilgilerinin yönetildiği alandır.")
    add_steps(doc, [
        "Yeni personel ekle düğmesini kullanın.",
        "Ad soyad, kurumsal e-posta, sicil numarası ve telefon bilgisini girin.",
        "İşyeri, bölüm ve rolü seçin.",
        "Kaydı oluşturduktan sonra hesap durumunu ve atamaları listeden doğrulayın.",
    ])
    add_image(doc, "10-personel.png", "Şekil 10 — Personel listesi ve yeni personel formu", 5.65)
    add_callout(doc, "Yetkilendirme", "Yönetici rolü yalnızca görev gereği ihtiyaç duyan kişilere verilmelidir. Personel rolü, kullanıcının yalnızca kendi verilerini görmesini sağlar.")

    new_page(doc, "5 · Raporlama", "Raporlama Merkezi", "Raporlama Merkezi, puantaj sonuçları ile sistemde gerçekleşen yönetim işlemlerini ayrı başlıklar altında sunar.")
    add_image(doc, "11-raporlama-merkezi.png", "Şekil 11 — Raporlama Merkezi", 6.2)
    add_bullets(doc, [
        "Puantaj Raporu: personel bazlı günlük çalışma sonuçları",
        "Giriş-Çıkış Raporu: ham geçiş hareketleri",
        "Denetim Kayıtları: kritik işlem geçmişi",
        "Düzeltme Geçmişi: puantaj düzeltme talepleri ve kararları",
    ])

    new_page(doc, "5.1 · Raporlama", "PDKS Raporları", "Tarih ve personel filtreleriyle günlük çalışma süresi, gecikme, erken çıkış, eksik süre ve fazla mesai sonuçlarını inceleyin.")
    add_steps(doc, [
        "Başlangıç ve bitiş tarihini seçin.",
        "Gerekirse personel filtresi uygulayın.",
        "Filtrele düğmesine basın ve sonuçları kontrol edin.",
        "Dışa aktar seçeneğiyle raporu CSV olarak indirin.",
    ])
    add_image(doc, "12-pdks-raporlari.png", "Şekil 12 — Tarih filtreli puantaj raporu", 5.55)

    new_page(doc, "5.2 · İzlenebilirlik", "Denetim Kayıtları", "Kritik yönetim işlemlerinin kim tarafından, ne zaman ve hangi kayıt üzerinde yapıldığını buradan izleyin.")
    add_image(doc, "13-denetim-kayitlari.png", "Şekil 13 — Denetim kayıtları", 5.55)
    add_bullets(doc, [
        "Tarih aralığı ve işlem türüyle arama yapın.",
        "Kullanıcı ve hedef kayıt bilgilerini birlikte değerlendirin.",
        "Eski ve yeni değerleri karşılaştırarak değişikliğin etkisini doğrulayın.",
    ])

    new_page(doc, "5.3 · İşlemler", "Onay ve Düzeltme İşlemleri", "İşlemler Merkezi, sonuçlandırma bekleyen kayıtları ve geçmiş işlem listelerini tek noktada toplar.")
    add_image(doc, "14-islemler-merkezi.png", "Şekil 14 — İşlemler Merkezi", 6.15, 2.65)
    add_steps(doc, [
        "Bekleyen İşler alanından izin veya puantaj düzeltme listesine geçin.",
        "Puantaj düzeltmesinde personelin önerdiği giriş-çıkış saatlerini ve gerekçeyi inceleyin.",
        "Vardiya ve geçiş kayıtlarıyla karşılaştırdıktan sonra talebi onaylayın ya da reddedin.",
    ])
    add_image(doc, "15-puantaj-duzeltmeleri.png", "Şekil 15 — Puantaj düzeltme talepleri", 6.1, 2.3)

    new_page(doc, "6 · Hesap", "Hesabım ve Genel Araçlar", "Sağ üstteki Hesabım alanından kişisel bilgiler, görünüm, güvenlik, bildirim ve kişisel veri seçeneklerine erişilir.")
    add_image(doc, "16-hesabim.png", "Şekil 16 — Kişisel bilgiler ekranı", 6.2)
    add_bullets(doc, [
        "Kişisel bilgilerim: ad, telefon ve profil fotoğrafı",
        "Arayüz ayarlarım: açık/koyu tema ve Türkçe/İngilizce dil",
        "Güvenlik: parola işlemleri",
        "Bildirimler: e-posta ve SMS tercihleri",
        "Verilerim: kişisel verilerin JSON olarak dışa aktarılması",
        "Genel arama: yetkiniz dahilindeki personel ve sayfalara hızlı erişim",
    ])

    new_page(doc, "7 · Personel Portalı", "Personel Kontrol Paneli", "Personel hesabı açıldığında yalnızca kullanıcının kendi çalışma ve izin bilgilerini içeren sadeleştirilmiş portal görüntülenir.")
    add_bullets(doc, [
        "Bugünkü durum ve çalışılan süre",
        "Aylık fazla mesai ve eksik kayıt özeti",
        "İlk giriş ve son çıkış bilgisi",
        "Bekleyen izin talepleri ve bildirimler",
    ])

    new_page(doc, "7.1 · Personel Portalı", "Çalışma Konumlarım ve Kayıtlarım", "Personel, kendisine atanmış konumları görür; gerektiğinde saha çalışması talebi oluşturur ve kendi puantaj kayıtlarını inceler.")
    add_image(doc, "18-calisma-konumlarim.png", "Şekil 17 — Personelin çalışma konumları ve talepleri", 6.05, 3.5)
    add_steps(doc, [
        "Çalışma Konumlarım ekranında aktif atamalarınızı kontrol edin.",
        "Yeni bir saha/konum ihtiyacı varsa tarih, konum türü ve gerekçeyle talep oluşturun.",
        "Çalışma Kayıtlarım ekranında tarih aralığını seçip günlük sonuçları inceleyin.",
        "Eksik veya hatalı bir gün için düzeltme talebi oluşturun.",
    ])

    new_page(doc, "7.2 · Personel Portalı", "İzinlerim ve Bildirimlerim", "Personel yeni izin talebi oluşturabilir, geçmiş taleplerinin durumunu ve yönetici kararlarını izleyebilir.")
    add_steps(doc, [
        "İzinlerim ekranında izin türü, gün bölümü ve tarih aralığını seçin.",
        "Gerekli açıklamayı yazıp talebi gönderin.",
        "Talep durumunu aynı ekrandan izleyin.",
        "Bildirimlerim alanından onay, ret ve puantaj uyarılarını takip edin.",
    ])
    add_image(doc, "21-bildirimlerim.png", "Şekil 18 — Bildirim tercihleri ve bildirim geçmişi", 6.0, 4.2)

    new_page(doc, "8 · Destek", "Sık Karşılaşılan Durumlar", "Aşağıdaki kontroller, en sık görülen kullanım sorunlarının hızlıca ayrıştırılmasına yardımcı olur.")
    issues = [
        ("Giriş yapamıyorum", "E-posta ve parolayı kontrol edin; hesabın aktif ve onaylı olduğunu yöneticinize doğrulatın."),
        ("QR okutulamıyor", "QR kodunun aktif ve güncel olduğunu, doğru giriş/çıkış türünde üretildiğini kontrol edin."),
        ("Puantaj sonucu yanlış", "Vardiya ataması, çalışma takvimi ve giriş-çıkış kayıtlarını birlikte inceleyin; gerekiyorsa düzeltme talebi oluşturun."),
        ("İzin görünmüyor", "Talebin onay durumunu ve tarih aralığını kontrol edin; gün bölümü seçiminin doğru olduğundan emin olun."),
        ("Bildirim gelmiyor", "Hesabım > Bildirimler bölümünde kanal tercihlerini ve iletişim bilgilerinizi kontrol edin."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2500, 6860])
    shade_cell(table.rows[0].cells[0], PURPLE)
    shade_cell(table.rows[0].cells[1], PURPLE)
    set_font(table.rows[0].cells[0].paragraphs[0].add_run("Durum"), bold=True, color=WHITE)
    set_font(table.rows[0].cells[1].paragraphs[0].add_run("Kontrol"), bold=True, color=WHITE)
    for title, answer in issues:
        cells = table.add_row().cells
        set_font(cells[0].paragraphs[0].add_run(title), bold=True, color=NAVY)
        set_font(cells[1].paragraphs[0].add_run(answer), color=NAVY)
        set_cell_margins(cells[0]); set_cell_margins(cells[1])
    set_table_geometry(table, [2500, 6860])
    doc.add_paragraph()
    add_callout(doc, "Güvenli kullanım", "Parolanızı paylaşmayın; ortak cihazlarda Oturumu açık tut seçeneğini kullanmayın. İşiniz bittiğinde Çıkış düğmesiyle oturumu kapatın. Kişisel verileri yalnızca görev ve yetki kapsamında görüntüleyin.", fill="ECFDF5", accent=GREEN)

    props = doc.core_properties
    props.title = "Faydam PDKS Kullanım ve Yardım Dokümanı"
    props.subject = "Web yönetim paneli ve personel portalı kullanım kılavuzu"
    props.author = "Faydam PDKS"
    props.keywords = "Faydam, PDKS, kullanım kılavuzu, yardım, puantaj, vardiya, izin"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
