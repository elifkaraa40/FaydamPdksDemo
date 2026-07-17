from pathlib import Path
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "MOBIL_UYGULAMA_ENTEGRASYON_REHBERI.md"
OUTPUT = ROOT / "docs" / "Faydam_PDKS_Mobil_Uygulama_Entegrasyon_Rehberi.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
GRAY = "5B6573"


def set_font(run, name="Calibri", size=None, bold=None, color=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Sayfa ")
    set_font(run, size=9, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in [s.name for s in doc.styles]:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        code.font.size = Pt(8.5)
        code.paragraph_format.left_indent = Inches(0.18)
        code.paragraph_format.right_indent = Inches(0.18)
        code.paragraph_format.space_before = Pt(4)
        code.paragraph_format.space_after = Pt(7)
        code.paragraph_format.line_spacing = 1.0

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hr = header.add_run("FAYDAM PDKS  |  MOBİL ENTEGRASYON")
    set_font(hr, size=8.5, bold=True, color=GRAY)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(72)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("TEKNİK ENTEGRASYON REHBERİ")
    set_font(r, size=11, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Faydam PDKS\nMobil Uygulama")
    set_font(r, size=28, bold=True, color=DARK_BLUE)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Kayıt, oturum, QR, mola, izin, puantaj ve bildirim sözleşmesi")
    set_font(r, size=13, color=GRAY)
    p.paragraph_format.space_after = Pt(54)

    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    widths = [Inches(1.875), Inches(4.625)]
    values = [
        ("Hedef ekip", "Flutter mobil uygulama geliştiricisi"),
        ("Durum", "Backend sözleşmesi ve kabul kriterleri kesinleşmiştir"),
        ("Son güncelleme", "17 Temmuz 2026"),
    ]
    for row, (label, value) in zip(table.rows, values):
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            set_cell_margins(cell)
        lr = row.cells[0].paragraphs[0].add_run(label)
        set_font(lr, bold=True, color=DARK_BLUE)
        vr = row.cells[1].paragraphs[0].add_run(value)
        set_font(vr)
        row.cells[0]._tc.get_or_add_tcPr().append(OxmlElement("w:shd"))
        row.cells[0]._tc.tcPr[-1].set(qn("w:fill"), LIGHT_BLUE)
    header_row_pr = table.rows[0]._tr.get_or_add_trPr()
    header_marker = OxmlElement("w:tblHeader")
    header_marker.set(qn("w:val"), "true")
    header_row_pr.append(header_marker)
    doc.add_page_break()


def add_inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, name="Consolas", size=9, color=DARK_BLUE)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_font(run, bold=True)
        else:
            paragraph.add_run(part)


def add_markdown(doc, source):
    lines = source.read_text(encoding="utf-8").splitlines()
    in_code = False
    code_lines = []
    first_h1_skipped = False
    for line in lines:
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="Code Block")
                shade_paragraph(p, "F4F6F9")
                run = p.add_run("\n".join(code_lines))
                set_font(run, name="Consolas", size=8.5)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("# ") and not first_h1_skipped:
            first_h1_skipped = True
            continue
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
        elif line.startswith("# "):
            doc.add_paragraph(line[2:].strip(), style="Heading 1")
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        elif line.startswith("- [ ] "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, "☐ " + line[6:])
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        else:
            p = doc.add_paragraph()
            add_inline(p, line)


def main():
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_markdown(doc, SOURCE)
    props = doc.core_properties
    props.title = "Faydam PDKS Mobil Uygulama Entegrasyon Rehberi"
    props.subject = "Mobil geliştirici için API ve kabul kriterleri"
    props.author = "Faydam PDKS"
    props.keywords = "PDKS, mobil, Flutter, API, QR, izin, mola"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
